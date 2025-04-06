from docx import Document

# Crear el documento
doc = Document()

# Título principal
doc.add_heading('Introducción a HTML', level=1)

# ¿Qué es HTML?
doc.add_heading('¿Qué es HTML?', level=2)
doc.add_paragraph(
    "HTML (HyperText Markup Language) es el lenguaje estándar para la creación de páginas web. "
    "No es un lenguaje de programación, sino un lenguaje de marcado que se usa para estructurar "
    "el contenido de una página web."
)

# ¿Por qué es importante?
doc.add_heading('¿Por qué es importante?', level=2)
doc.add_paragraph("• Es el fundamento de todas las páginas web.")
doc.add_paragraph("• Permite organizar el contenido en elementos como encabezados, párrafos, enlaces e imágenes.")
doc.add_paragraph("• Funciona junto con CSS y JavaScript para dar estilo e interactividad a las páginas.")

# Estructura básica de un documento HTML
doc.add_heading('Estructura básica de un documento HTML', level=2)
doc.add_paragraph(
    "Un archivo HTML tiene una estructura base que define el contenido de una página web. "
    "A continuación, un ejemplo de una página HTML mínima:"
)

# Código HTML de ejemplo
codigo_html = """
<!DOCTYPE html>
<html>
<head>
    <title>Mi primera página web</title>
</head>
<body>
    <h1>¡Hola, mundo!</h1>
    <p>Este es un párrafo de ejemplo.</p>
</body>
</html>
"""
doc.add_paragraph(codigo_html, style='Normal')

# Explicación del código
doc.add_heading('Explicación del código:', level=3)
doc.add_paragraph("1. <!DOCTYPE html> - Indica que el documento es un HTML5.")
doc.add_paragraph("2. <html> - Contenedor principal de la página.")
doc.add_paragraph("3. <head> - Contiene información sobre la página, como el título.")
doc.add_paragraph("4. <title> - Define el título que aparece en la pestaña del navegador.")
doc.add_paragraph("5. <body> - Contiene el contenido visible de la página.")
doc.add_paragraph("6. <h1> - Es un encabezado de nivel 1 (título principal).")
doc.add_paragraph("7. <p> - Representa un párrafo de texto.")

# Estructura de etiquetas en HTML
doc.add_heading('Estructura de las etiquetas en HTML', level=2)
doc.add_paragraph(
    "En HTML, las etiquetas generalmente tienen una estructura de apertura y cierre:\n\n"
    "    <etiqueta>Contenido</etiqueta>\n\n"
    "Por ejemplo:\n\n"
    "    <p>Este es un párrafo</p>\n\n"
    "Algunas etiquetas no tienen cierre y se llaman etiquetas vacías, como <img> y <br>:\n\n"
    "    <img src='imagen.jpg' alt='Descripción de la imagen'>\n"
    "    <br> <!-- Esto es un salto de línea -->"
)

# Uso de comentarios en HTML
doc.add_heading('Uso de comentarios en HTML', level=2)
doc.add_paragraph(
    "Los comentarios se utilizan para hacer anotaciones en el código sin que sean visibles en la página. "
    "Son útiles para dejar recordatorios o explicaciones dentro del código.\n\n"
    "    <!-- Esto es un comentario en HTML -->\n"
    "    <p>Este es un párrafo visible</p>\n\n"
    "Los navegadores ignoran los comentarios, por lo que no afectan la visualización de la página."
)

# Guardar el documento
doc_filename = "C:/Users/miguel/Desktop/Mis Documentos/talento tech/frontEnd_JS/clases/clase 1/Introduccion_HTML.docx"

doc.save(doc_filename)

# Devolver el archivo al usuario
doc_filename
